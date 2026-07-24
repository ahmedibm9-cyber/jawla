from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ---- Page Setup ----
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Tahoma'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper: set cell shading
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper: set RTL on cell
def set_cell_rtl(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    rtl = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
    tcPr.append(rtl)

# Helper: add RTL paragraph
def add_rtl_paragraph(doc, text, bold=False, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.RIGHT, color=RGBColor(0x1A,0x1A,0x1A), space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Tahoma'
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    return p

# ============================================================
# HEADER WATERMARK
# ============================================================
p = add_rtl_paragraph(doc, 'مسودة — غير ملزم إلا بالتوقيع', bold=True, size=Pt(13), color=RGBColor(0x99,0x99,0x99), space_after=Pt(20))
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="dashed" w:sz="4" w:space="1" w:color="999999"/><w:top w:val="dashed" w:sz="4" w:space="1" w:color="999999"/></w:pBdr>')
pPr.append(pBdr)

# ============================================================
# TITLE
# ============================================================
add_rtl_paragraph(doc, 'ملخص الصفقة التجارية', bold=True, size=Pt(22), color=RGBColor(0x2C,0x3E,0x50), space_after=Pt(2))
add_rtl_paragraph(doc, 'JAWLA — ترخيص برمجيات + صيانة + استضافة', bold=False, size=Pt(14), color=RGBColor(0x7F,0x8C,0x8D), space_after=Pt(20))

# Horizontal rule
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="2C3E50"/></w:pBdr>')
pPr.append(pBdr)

# ============================================================
# SECTION 1: البيانات الأساسية
# ============================================================
add_rtl_paragraph(doc, '١. البيانات الأساسية', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(15), space_after=Pt(10))

fields = [
    'رقم الصفقة', 'التاريخ', 'اسم العميل (الكيان القانوني)', 'عنوان العميل',
    'البريد الإلكتروني للعميل', 'هاتف العميل', 'اسم المورد (الكيان القانوني)',
    'عنوان المورد', 'البريد الإلكتروني للمورد', 'هاتف المورد'
]

table = doc.add_table(rows=len(fields), cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, field in enumerate(fields):
    row = table.rows[i]
    cell_l = row.cells[0]
    cell_l.text = ''
    p = cell_l.paragraphs[0]
    run = p.add_run(field)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(cell_l, 'F8F9FA')
    set_cell_rtl(cell_l)
    cell_r = row.cells[1]
    cell_r.text = ''
    set_cell_rtl(cell_r)

for row in table.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(12)

# ============================================================
# SECTION 2: نطاق الترخيص
# ============================================================
add_rtl_paragraph(doc, '٢. نطاق الترخيص', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(20), space_after=Pt(10))

license_data = [
    ('نوع الترخيص', 'دائم (Perpetual) — غير محدود المدة'),
    ('الإقليم المغطى', 'مصر + المملكة العربية السعودية فقط'),
    ('عدد المستخدمين', 'غير محدود (Enterprise / Unlimited Seats)'),
    ('نموذج النشر', 'استضافة المورد (SaaS Single-Tenant) — العميل يدفع تكلفة الاستضافة'),
]

table = doc.add_table(rows=len(license_data)+1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0]
for j, txt in enumerate(['البند', 'التفاصيل']):
    cell = hdr.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(cell, '2C3E50')
    set_cell_rtl(cell)

for i, (label, val) in enumerate(license_data):
    row = table.rows[i+1]
    for j, txt in enumerate([label, val]):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Tahoma'
        run.font.size = Pt(10)
        run.font.bold = (j == 0)
        run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_rtl(cell)
        if i % 2 == 0:
            set_cell_shading(cell, 'F8F9FA')

# ============================================================
# SECTION 3: الرسوم
# ============================================================
add_rtl_paragraph(doc, '٣. الرسوم (بالجنيه المصري — ج.م)', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(20), space_after=Pt(10))

fees = [
    ('رسوم الترخيص (License Fee)', '', 'مرة واحدة', ''),
    ('رسوم الصيانة السنوية (Annual Maintenance)', '', 'سنوياً مقدماً', ''),
    ('تكلفة الاستضافة التقديرية (Hosting Estimate)', '', 'شهرياً آجلاً', 'تمرير التكلفة الفعلية (VPS + Cloudflare)'),
    ('الخدمات المهنية (Professional Services)', '', 'وفق مراحل (SOW)', ''),
    ('أخرى', '', '', ''),
]

table = doc.add_table(rows=len(fees)+2, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['البند', 'المبلغ (ج.م)', 'التكرار', 'ملاحظات']
hdr = table.rows[0]
for j, txt in enumerate(headers):
    cell = hdr.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [1,2] else WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(cell, '2C3E50')
    set_cell_rtl(cell)

for i, (item, amt, freq, notes) in enumerate(fees):
    row = table.rows[i+1]
    for j, txt in enumerate([item, amt, freq, notes]):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Tahoma'
        run.font.size = Pt(10)
        run.font.bold = (j == 0)
        run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [1,2] else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_rtl(cell)
        if i % 2 == 0:
            set_cell_shading(cell, 'F8F9FA')

# Total row
total_row = table.rows[-1]
for j, txt in enumerate(['الإجمالي التقديري للسنة الأولى', '', '—', '']):
    cell = total_row.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [1,2] else WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(cell, 'E8F5E9')
    set_cell_rtl(cell)

# ============================================================
# SECTION 4: جدول الدفعات
# ============================================================
add_rtl_paragraph(doc, '٤. جدول الدفعات (Milestone-Based)', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(20), space_after=Pt(10))

milestones = [
    'عند التوقيع (Signing)',
    'الإعداد والتكوين (Setup & Configuration)',
    'اختبار قبول المستخدم (UAT)',
    'الإطلاق (Go-Live)',
    'أخرى / دفعة نهائية',
]

table = doc.add_table(rows=len(milestones)+1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['المرحلة', 'النسبة %', 'المبلغ (ج.م)', 'تاريخ الاستحقاق', 'حالة الدفع']
hdr = table.rows[0]
for j, txt in enumerate(headers):
    cell = hdr.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '2C3E50')
    set_cell_rtl(cell)

for i, ms in enumerate(milestones):
    row = table.rows[i+1]
    for j, txt in enumerate([ms, '', '', '', '☐ مدفوع']):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Tahoma'
        run.font.size = Pt(10)
        run.font.bold = (j == 0)
        run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_rtl(cell)
        if i % 2 == 0:
            set_cell_shading(cell, 'F8F9FA')

# ============================================================
# SECTION 5: مستوى الدعم
# ============================================================
add_rtl_paragraph(doc, '٥. مستوى الدعم (Support Tier)', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(20), space_after=Pt(10))

table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['المستوى', 'القنوات', 'زمن الاستجابة (هدف)', 'ساعات العمل']
hdr = table.rows[0]
for j, txt in enumerate(headers):
    cell = hdr.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(cell, '2C3E50')
    set_cell_rtl(cell)

data = ['قياسي (Standard) — مشمول', 'البريد الإلكتروني + نظام التذاكر',
        'حرج (P1): 4 ساعات عمل\nعاجل (P2): يوم عمل واحد\nعادي (P3): 3 أيام عمل',
        'الأحد–الخميس، ٩ص–٥م (توقيت القاهرة)']
row = table.rows[1]
for j, txt in enumerate(data):
    cell = row.cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Tahoma'
    run.font.size = Pt(10)
    run.font.bold = (j == 0)
    run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_rtl(cell)

# Note
p = add_rtl_paragraph(doc, 'ملاحظة: لا يوجد SLA لوقت التشغيل (Uptime SLA) — جهد معقول (Best Effort). راجع اتفاقية المستوى الرئيسي للتفاصيل الكاملة.',
    size=Pt(9), color=RGBColor(0x85,0x64,0x04), space_before=Pt(8), space_after=Pt(15))
pPr = p._p.get_or_add_pPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3CD" w:val="clear"/>')
pPr.append(shd)

# ============================================================
# SECTION 6: مدة الصيانة والتجديد
# ============================================================
add_rtl_paragraph(doc, '٦. مدة الصيانة والتجديد', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(15), space_after=Pt(10))

maint_data = [
    ('دورة الصيانة', 'سنة واحدة (تجديد تلقائي)'),
    ('إشعار عدم التجديد', '٦٠ يوماً قبل تاريخ التجديد (كتابةً)'),
    ('تاريخ بداية الصيانة', 'تاريخ الإطلاق (Go-Live) — ______________________'),
    ('تاريخ التجديد القادم', '______________________'),
]

table = doc.add_table(rows=len(maint_data), cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, val) in enumerate(maint_data):
    row = table.rows[i]
    for j, txt in enumerate([label, val]):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Tahoma'
        run.font.size = Pt(10)
        run.font.bold = (j == 0)
        run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_rtl(cell)
        if j == 0:
            set_cell_shading(cell, 'F8F9FA')

# ============================================================
# SECTION 7: بنود خاصة / ملاحظات
# ============================================================
add_rtl_paragraph(doc, '٧. بنود خاصة / ملاحظات إضافية', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(20), space_after=Pt(10))

table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]
cell.text = ''
for _ in range(6):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
set_cell_rtl(cell)

# ============================================================
# SECTION 8: التوقيعات
# ============================================================
add_rtl_paragraph(doc, '٨. التوقيعات', bold=True, size=Pt(14), color=RGBColor(0x2C,0x3E,0x50), space_before=Pt(25), space_after=Pt(15))

table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, title in enumerate(['المورد (Provider)', 'العميل (Customer)']):
    cell = table.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Tahoma'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    set_cell_shading(cell, '2C3E50')
    set_cell_rtl(cell)

sig_fields = ['الاسم:', 'الصفة/المنصب:', 'التاريخ:', 'التوقيع:']
for j in range(2):
    cell = table.rows[1].cells[j]
    cell.text = ''
    for field in sig_fields:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(field + ' ___________________________________')
        run.font.name = 'Tahoma'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    set_cell_rtl(cell)

# Footer note
add_rtl_paragraph(doc, 'هذه الوثيقة ملخص للصفقة فقط. الاتفاقية الملزمة هي "اتفاقية ترخيص برمجيات JAWLA" الموقعة من الطرفين.',
    size=Pt(9), color=RGBColor(0x99,0x99,0x99), space_before=Pt(30), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph(doc, 'JAWLA Software License Agreement — Deal Summary Template v1.0',
    size=Pt(8), color=RGBColor(0x99,0x99,0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Save
doc.save('Deal_Summary_Arabic.docx')
print('Deal_Summary_Arabic.docx created successfully')